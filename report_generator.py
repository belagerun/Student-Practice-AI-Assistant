import re
from datetime import datetime
from pathlib import Path


ARTIFACTS_DIR = Path(__file__).resolve().with_name("artifacts")


def _clean_text(text):
    return " ".join((text or "").split())


def clean_markdown_text(text):
    text = text or ""
    text = re.sub(r"!\[([^\]]*)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
    text = re.sub(r"(\*\*|__)(.*?)\1", r"\2", text)
    text = re.sub(r"(\*|_)(.*?)\1", r"\2", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"^\s*#{1,6}\s*", "", text, flags=re.MULTILINE)
    text = re.sub(r"^\s*[-*+]\s+", "", text, flags=re.MULTILINE)

    return _clean_text(text)


def _set_run_style(run, size_pt=12, bold=False):
    from docx.shared import Pt, RGBColor

    run.bold = bold
    run.underline = False
    run.font.name = "Times New Roman"
    run.font.size = Pt(size_pt)
    run.font.color.rgb = RGBColor(0, 0, 0)


def _configure_paragraph(paragraph, justify=True, first_line=True):
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Inches

    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.JUSTIFY
        if justify
        else WD_ALIGN_PARAGRAPH.LEFT
    )
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.first_line_indent = (
        Inches(0.5)
        if first_line
        else None
    )
    paragraph.paragraph_format.space_after = 0


def _configure_report_styles(document):
    from docx.enum.style import WD_STYLE_TYPE
    from docx.shared import Pt, RGBColor

    normal_style = document.styles["Normal"]
    normal_style.font.name = "Times New Roman"
    normal_style.font.size = Pt(12)
    normal_style.font.color.rgb = RGBColor(0, 0, 0)
    normal_style.paragraph_format.line_spacing = 1.15

    custom_styles = [
        ("CustomTitle", 18),
        ("CustomHeading1", 14),
        ("CustomHeading2", 13),
    ]

    for style_name, size_pt in custom_styles:
        try:
            style = document.styles[style_name]
        except KeyError:
            style = document.styles.add_style(
                style_name,
                WD_STYLE_TYPE.PARAGRAPH
            )

        style.font.name = "Times New Roman"
        style.font.size = Pt(size_pt)
        style.font.bold = True
        style.font.underline = False
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.line_spacing = 1.15
        style.paragraph_format.first_line_indent = None


def add_black_paragraph(document, text):
    paragraph = document.add_paragraph(style="Normal")
    _configure_paragraph(paragraph, justify=True, first_line=True)
    run = paragraph.add_run(clean_markdown_text(text))
    _set_run_style(run, 12, False)

    return paragraph


def add_black_heading(document, text, level):
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    style_name = "CustomHeading1"
    size_pt = 14

    if level == 0:
        style_name = "CustomTitle"
        size_pt = 18
    elif level == 2:
        style_name = "CustomHeading2"
        size_pt = 13

    paragraph = document.add_paragraph(style=style_name)
    paragraph.alignment = (
        WD_ALIGN_PARAGRAPH.CENTER
        if level == 0
        else WD_ALIGN_PARAGRAPH.LEFT
    )
    paragraph.paragraph_format.line_spacing = 1.15
    paragraph.paragraph_format.first_line_indent = None
    paragraph.paragraph_format.space_after = 0
    run = paragraph.add_run(clean_markdown_text(text))
    _set_run_style(run, size_pt, True)

    return paragraph


def _split_sections(content):
    content = (content or "").strip()

    if not content:
        raise ValueError("Report content is empty.")

    sections = []
    current_title = ""
    current_lines = []

    for raw_line in content.splitlines():
        line = raw_line.strip()

        if not line:
            continue

        heading_match = re.match(
            r"^(?:#{1,3}\s*)?(?:section\s*\d+|раздел\s*\d+)?\s*[:.-]?\s*(.+)$",
            line,
            flags=re.IGNORECASE,
        )

        looks_like_heading = (
            len(line) <= 90
            and not line.startswith(("-", "*"))
            and (
                line.endswith(":")
                or re.match(r"^(?:#{1,3}|section|раздел|\d+[.)])", line, re.IGNORECASE)
            )
        )

        if looks_like_heading and heading_match:
            if current_title or current_lines:
                sections.append(
                    {
                        "title": _clean_text(current_title) or "Section",
                        "body": _clean_text(" ".join(current_lines)),
                    }
                )

            current_title = clean_markdown_text(
                heading_match.group(1).strip(" #:.-")
            )
            current_lines = []
            continue

        current_lines.append(clean_markdown_text(line.strip(" -*\t")))

    if current_title or current_lines:
        sections.append(
            {
                "title": _clean_text(current_title) or "Main Part",
                "body": _clean_text(" ".join(current_lines)),
            }
        )

    if len(sections) >= 2:
        return sections

    paragraphs = [
        clean_markdown_text(paragraph)
        for paragraph in re.split(r"\n{2,}", content)
        if paragraph.strip()
    ]

    if paragraphs:
        return [
            {
                "title": f"Section {index}",
                "body": paragraph,
            }
            for index, paragraph in enumerate(paragraphs, start=1)
        ]

    return [{"title": "Main Part", "body": _clean_text(content)}]


def _add_paragraphs(document, text):
    sentences = re.split(r"(?<=[.!?])\s+", clean_markdown_text(text))
    current = []

    for sentence in sentences:
        if not sentence:
            continue

        current.append(sentence)

        if len(" ".join(current)) >= 450:
            add_black_paragraph(document, " ".join(current))
            current = []

    if current:
        add_black_paragraph(document, " ".join(current))


def generate_docx_report(title, content, output_path, metadata=None):
    try:
        from docx import Document
        from docx.enum.text import WD_ALIGN_PARAGRAPH
    except ImportError as error:
        raise RuntimeError(
            "python-docx is not installed. Add python-docx to requirements.txt."
        ) from error

    metadata = metadata or {}
    title = _clean_text(title) or "PracticeAI Report"
    sections = _split_sections(content)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    document = Document()
    _configure_report_styles(document)

    title_paragraph = add_black_heading(
        document,
        title,
        0
    )
    title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    subtitle = document.add_paragraph(style="Normal")
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle.paragraph_format.first_line_indent = None
    subtitle.paragraph_format.line_spacing = 1.15
    subtitle_run = subtitle.add_run(
        f"Created: {datetime.now().strftime('%Y-%m-%d %H:%M')}"
    )
    _set_run_style(subtitle_run, 12, False)

    generated = document.add_paragraph(style="Normal")
    generated.alignment = WD_ALIGN_PARAGRAPH.CENTER
    generated.paragraph_format.first_line_indent = None
    generated.paragraph_format.line_spacing = 1.15
    generated_run = generated.add_run("Generated by PracticeAI")
    _set_run_style(generated_run, 12, False)

    document.add_page_break()

    intro_text = metadata.get("introduction") or (
        "This report summarizes the key information, analysis, and conclusions "
        "prepared from the provided request and available context."
    )
    conclusion_text = metadata.get("conclusion") or (
        "The report highlights the main findings and provides a concise basis "
        "for further discussion, study, or decision-making."
    )

    add_black_heading(document, "Introduction", 1)
    _add_paragraphs(document, intro_text)

    add_black_heading(document, "Main Part", 1)

    for section in sections:
        section_title = section.get("title") or "Section"
        section_body = section.get("body") or ""

        if not section_body:
            continue

        add_black_heading(document, section_title[:90], 2)
        _add_paragraphs(document, section_body)

    add_black_heading(document, "Conclusion", 1)
    _add_paragraphs(document, conclusion_text)

    sources = metadata.get("sources") or []

    if sources:
        add_black_heading(document, "Used Documents", 1)

        for source in sources:
            add_black_paragraph(document, f"- {source}")

    add_black_paragraph(document, "Generated by PracticeAI")
    document.save(output_path)

    if not output_path.exists() or output_path.stat().st_size == 0:
        raise RuntimeError("DOCX file was not saved correctly.")

    return {
        "file_path": str(output_path),
        "sections_count": len([section for section in sections if section.get("body")]),
        "file_size": output_path.stat().st_size,
    }


def make_report_output_path():
    ARTIFACTS_DIR.mkdir(parents=True, exist_ok=True)
    file_name = f"report_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
    return ARTIFACTS_DIR / file_name
